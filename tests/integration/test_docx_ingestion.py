from io import BytesIO
from uuid import uuid4

import pytest
import pytest_asyncio
from docx import Document as WordDocument
from sqlalchemy import delete, select

from app.db.session import AsyncSessionLocal, engine
from app.ingestion.service import IngestionService
from app.models.chunk import Chunk
from app.models.document import Document
from app.models.workspace import Workspace


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@pytest_asyncio.fixture(autouse=True)
async def isolate_engine_pool():
    await engine.dispose()
    yield
    await engine.dispose()


def make_docx(marker: str) -> bytes:
    document = WordDocument()
    document.add_heading("DOCX Integration", level=1)
    document.add_paragraph(f"Unique DOCX marker: {marker}")
    document.add_heading("Controls", level=2)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Control"
    table.cell(0, 1).text = "Behavior"
    table.cell(1, 0).text = "Workspace"
    table.cell(1, 1).text = "Fail closed"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.asyncio
async def test_docx_ingestion_persists_chunks_and_structure() -> None:
    marker = uuid4().hex
    file_bytes = make_docx(marker)
    workspace_id: int | None = None

    try:
        async with AsyncSessionLocal() as session:
            workspace = Workspace(name=f"docx-integration-{uuid4().hex}")
            session.add(workspace)
            await session.commit()
            await session.refresh(workspace)
            workspace_id = workspace.id

            result = await IngestionService(session=session).ingest(
                workspace_id=workspace.id,
                filename="system-design.docx",
                content_type=DOCX_MIME,
                file_bytes=file_bytes,
            )

            assert result.status == "COMPLETED"
            assert result.file_type == "docx"
            assert result.chunk_count >= 2

            stored_document = await session.get(Document, result.document_id)
            assert stored_document is not None
            assert stored_document.file_type == "docx"
            assert stored_document.content_type == DOCX_MIME

            chunks = (
                await session.scalars(
                    select(Chunk)
                    .where(Chunk.document_id == result.document_id)
                    .order_by(Chunk.chunk_index.asc())
                )
            ).all()

            assert any(marker in chunk.text for chunk in chunks)
            assert any(
                chunk.section == "DOCX Integration > Controls"
                for chunk in chunks
            )
            assert any(
                "table" in list(chunk.metadata_json.get("element_types", []))
                for chunk in chunks
            )
            assert any("Workspace | Fail closed" in chunk.text for chunk in chunks)
    finally:
        if workspace_id is not None:
            async with AsyncSessionLocal() as session:
                await session.execute(
                    delete(Workspace).where(Workspace.id == workspace_id)
                )
                await session.commit()
