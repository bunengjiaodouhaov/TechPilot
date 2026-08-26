from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document

from app.ingestion.chunker import StructureAwareChunker
from app.ingestion.parsers.docx import DOCXParser
from app.ingestion.schemas import ParseInput


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def make_docx() -> bytes:
    document = Document()
    document.add_heading("System Architecture", level=1)
    document.add_paragraph("TechPilot keeps retrieval evidence scoped to a workspace.")
    document.add_heading("Runtime", level=2)
    document.add_paragraph("The runtime validates tool permissions before execution.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Component"
    table.cell(0, 1).text = "Boundary"
    table.cell(1, 0).text = "ToolRuntime"
    table.cell(1, 1).text = "Permission enforcement"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def parse_docx(file_bytes: bytes):
    return DOCXParser().parse(
        ParseInput(
            filename="architecture.docx",
            content_type=DOCX_MIME,
            file_size=len(file_bytes),
            file_bytes=file_bytes,
        )
    )


def test_docx_parser_preserves_headings_paragraphs_and_tables() -> None:
    parsed = parse_docx(make_docx())

    assert parsed.file_type == "docx"
    assert parsed.title == "System Architecture"
    assert parsed.metadata["table_count"] == 1
    assert parsed.metadata["embedded_images_interpreted"] is False

    element_types = [element.element_type for element in parsed.elements]
    assert element_types == ["heading", "paragraph", "heading", "paragraph", "table"]

    runtime_paragraph = parsed.elements[3]
    assert runtime_paragraph.source_metadata["heading_path"] == [
        "System Architecture",
        "Runtime",
    ]

    table = parsed.elements[4]
    assert table.source_metadata["table_index"] == 1
    assert table.source_metadata["row_count"] == 2
    assert table.text.startswith("Table 1:")
    assert "Component | Boundary" in table.text
    assert "ToolRuntime | Permission enforcement" in table.text


def test_docx_chunks_inject_heading_path_and_keep_table_type() -> None:
    parsed = parse_docx(make_docx())

    chunks = StructureAwareChunker(max_chars=1200).chunk(parsed)

    assert chunks
    assert any(
        chunk.section == "System Architecture > Runtime"
        for chunk in chunks
    )
    table_chunks = [
        chunk
        for chunk in chunks
        if "table" in chunk.metadata["element_types"]
    ]
    assert len(table_chunks) == 1
    assert "ToolRuntime | Permission enforcement" in table_chunks[0].text
    assert table_chunks[0].section == "System Architecture > Runtime"


def test_docx_parser_rejects_non_docx_bytes() -> None:
    with pytest.raises(ValueError, match="invalid or unreadable"):
        parse_docx(b"not-docx!")


def test_docx_parser_rejects_suspicious_archive_compression_ratio() -> None:
    archive_bytes = BytesIO(make_docx())
    with ZipFile(archive_bytes, mode="a", compression=ZIP_DEFLATED) as archive:
        archive.writestr("word/media/repetitive.bin", b"0" * (2 * 1024 * 1024))

    with pytest.raises(ValueError, match="compression ratio"):
        parse_docx(archive_bytes.getvalue())
